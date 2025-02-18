from dotenv import load_dotenv
from docx import Document
from pathlib import Path
from src.llm.llm import get_llm
from src.service.resume_parser import ResumeParser
from src.service.emotion_recognition import EmotionRecognition
from src.utils.utils import (
    extract_audio,
    audio2text,
    sample_frames,
    parse_yaml_string,
)
from src.template.grading_prompt import (
    GRADE_RESPONSE_PROMPT,
    RANKING_AND_FEEDBACK_PROMPT,
)

# sample input values
from src.archive.sample_inputs import (
    VIDEO_PATH,
    RESUME_PATH,
    INTERVIEW_QUESTION,
    JOB_REQUIREMENTS,
)

load_dotenv()

BASE_DIR = Path(__file__).resolve().parent.parent

# customise this part
LLM_CONFIG_FILE = BASE_DIR / "configs/llm/openai-gpt-3.5-turbo.yaml"
# LLM_CONFIG_FILE = "./src/configs/llm/openai-gpt-4o-mini.yaml"
# LLM_CONFIG_FILE = "./src/configs/llm/nvidia-llama-3.1-nemotron-70b-instruct.yaml"

RESUME_PARSER_CONFIG_FILE = BASE_DIR / "configs/parser/llamaparse_en.yaml"
OUTPUT_AUDIO_FILE = BASE_DIR / "output/audio_output.wav"  # only supports .wav
OUTPUT_REPORT_FILE = BASE_DIR / "output/report.docx"

# init API keys as env variables
load_dotenv()

# init LLM & resume parser
llm = get_llm(str(LLM_CONFIG_FILE))
parser = ResumeParser(str(RESUME_PARSER_CONFIG_FILE))


# 1. extract audio from video
OUTPUT_AUDIO_FILE = extract_audio(VIDEO_PATH, str(OUTPUT_AUDIO_FILE))
# assert OUTPUT_AUDIO_FILE is not None, f"Audio extraction failed."

# 2. audio to text
audio_text = audio2text(OUTPUT_AUDIO_FILE)
print(audio_text)

# 3. extract frames form video
frames = sample_frames(VIDEO_PATH, sample_rate=8)
print(frames)

# 4. deepface extract emotions & compite confidence scores
emotions = EmotionRecognition.detect_face_emotions(frames)
emotions_dict = EmotionRecognition.process_emotions(emotions)
conf_score = emotions_dict["conf"]
print(emotions_dict)

# 5. llamaparse parse resume into MD
resume_md = parser.parse_resume_to_markdown(RESUME_PATH)
print(resume_md)

# 6. llm grade question response
formatted_grading_prompt = GRADE_RESPONSE_PROMPT.format(
    interview_question=INTERVIEW_QUESTION,
    conf_score=conf_score,
    response_text=audio_text,
)
grade = llm.complete(formatted_grading_prompt)
print(grade)

# 7. llm rank and output final feedback
formatted_ranking_prompt = RANKING_AND_FEEDBACK_PROMPT.format(
    job_requirements=JOB_REQUIREMENTS, interview_feedback=grade, resume_text=resume_md
)
rank_and_feedback = llm.complete(formatted_ranking_prompt)
print(rank_and_feedback)


# 8. save to .docx report
expected_keys = ["name", "score", "feedback"]
rank_and_feedback_dict = parse_yaml_string(
    yaml_string=rank_and_feedback, expected_keys=expected_keys, cleanup=True
)
print(rank_and_feedback_dict)

doc = Document()
doc.add_heading(f"{rank_and_feedback_dict['name']}", 0)
doc.add_heading(f"Overall Score: {rank_and_feedback_dict['score']}", 1)
doc.add_heading(f"Brief Overview", 1)
doc.add_paragraph(f"{rank_and_feedback_dict['feedback']}")

# Save the document
doc.save(str(OUTPUT_REPORT_FILE))
