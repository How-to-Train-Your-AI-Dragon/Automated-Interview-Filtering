import gradio as gr
import pandas as pd
import logging
from pathlib import Path
from docx import Document
from typing import Optional, List
from dataclasses import dataclass
from dotenv import load_dotenv
from src.archive.sample_inputs import INTERVIEW_QUESTION, JOB_REQUIREMENTS
from src.configs.database.firebase import write_user_data, read_all_users
from src.llm.llm import get_llm
from src.service.emotion_recognition import EmotionRecognition
from src.service.resume_parser import ResumeParser
from src.utils.utils import (
    parse_yaml_string,
    extract_audio,
    audio2text,
    sample_frames,
)
from src.template.grading_prompt import (
    GRADE_RESPONSE_PROMPT,
    RANKING_AND_FEEDBACK_PROMPT,
)

load_dotenv()
# ENVIRONMENT = os.getenv("ENVIRONMENT", "local")

# Define base paths dynamically
# if ENVIRONMENT == "local":
#     BASE_DIR = Path(__file__).resolve().parent.parent.parent.parent
# else:  # Assume hosted on Hugging Face Spaces
BASE_DIR = Path(".").resolve()

LLM_CONFIG_FILE = BASE_DIR / "configs/llm/openai-gpt-3.5-turbo.yaml"
RESUME_PARSER_CONFIG_FILE = BASE_DIR / "configs/parser/llamaparse_en.yaml"
OUTPUT_AUDIO_FILE_EMPTY = BASE_DIR / "output/audio_output.wav"
OUTPUT_REPORT_FILE_EMPTY = BASE_DIR / "output/report.docx"


@dataclass
class ProcessingResult:
    candidate_name: Optional[str] = None
    candidate_score: Optional[int] = None
    candidate_feedbacks: Optional[List[str]] = None
    feedback_md: Optional[str] = None
    interview_question: Optional[str] = None
    job_requirements: Optional[str] = None
    error_message: Optional[str] = None


class GradioInterface:
    VALID_VIDEO_EXTENSIONS = {".mp4", ".avi", ".mkv"}
    VALID_RESUME_EXTENSIONS = {".pdf"}

    def __init__(self):
        self.parser = None
        self.llm = None
        self.logger = None
        self.candidate_feedback = pd.DataFrame(columns=["Name", "Score", "Feedback"])
        self.setup_logging()
        self.initialize_services()

    def setup_logging(self):
        logging.basicConfig(
            level=logging.INFO,
            format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
        )
        self.logger = logging.getLogger(__name__)

    def initialize_services(self):
        try:
            self.llm = get_llm(str(LLM_CONFIG_FILE))
            self.parser = ResumeParser(str(RESUME_PARSER_CONFIG_FILE))
        except Exception as e:
            self.logger.error(f"Failed to initialize services: {str(e)}")
            raise

    def validate_inputs(
        self,
        video_path: Optional[str],
        resume_path: Optional[str],
        interview_questions: Optional[str],
        job_requirements: Optional[str],
    ) -> Optional[str]:
        if not video_path:
            return "Please upload an interview video."
        if not resume_path:
            return "Please upload a resume (PDF)."
        if not interview_questions:
            return "Please provide interview questions."
        if not job_requirements:
            return "Please provide job requirements."
        if not self._validate_file_format(video_path, self.VALID_VIDEO_EXTENSIONS):
            return "Invalid video format."
        if not self._validate_file_format(resume_path, self.VALID_RESUME_EXTENSIONS):
            return "Please submit resume in PDF format."
        return None

    def _validate_file_format(self, file_path: str, valid_extensions: set) -> bool:
        return isinstance(file_path, str) and any(
            file_path.lower().endswith(ext) for ext in valid_extensions
        )

    def process_video(self, video_path: str) -> Optional[str]:
        OUTPUT_AUDIO_FILE = extract_audio(video_path, str(OUTPUT_AUDIO_FILE_EMPTY))
        audio_text = audio2text(OUTPUT_AUDIO_FILE)
        return audio_text

    def analyze_emotions(self, video_path: str) -> Optional[str]:
        frames = sample_frames(video_path, sample_rate=8)
        emotions = EmotionRecognition.detect_face_emotions(frames)
        emotions_dict = EmotionRecognition.process_emotions(emotions)
        conf_score = emotions_dict["conf"]
        return conf_score

    def process_resume(self, resume_path: str) -> Optional[str]:
        resume_md = self.parser.parse_resume_to_markdown(resume_path)
        return resume_md

    def format_feedback_to_markdown(self, feedback_df: pd.DataFrame) -> str:
        if feedback_df.empty:
            return "No feedback available."

        name = feedback_df["Name"].iloc[0]
        score = feedback_df["Score"].iloc[0]

        # Start with header
        markdown_text = f"""
# Candidate Assessment Report 📝

## Candidate Name ✨
{name}

## Candidate Overall Score  🎯
{score}/100

## Detailed Feedback  🛠️
"""

        for idx, row in feedback_df.iterrows():
            markdown_text += f"- {row['Feedback']}\n\n"

        return markdown_text

    def get_feedback(
        self,
        itv_question: str,
        job_requirements: str,
        conf_score: str,
        audio_text: str,
        resume_md: str,
    ) -> pd.DataFrame:

        formatted_grading_prompt = GRADE_RESPONSE_PROMPT.format(
            interview_question=itv_question,
            conf_score=conf_score,
            response_text=audio_text,
        )

        grade = self.llm.complete(formatted_grading_prompt)

        formatted_ranking_prompt = RANKING_AND_FEEDBACK_PROMPT.format(
            job_requirements=job_requirements,
            interview_feedback=grade,
            resume_text=resume_md,
        )
        rank_and_feedback = self.llm.complete(formatted_ranking_prompt)

        expected_keys = ["name", "score", "feedback"]
        rank_and_feedback_dict = parse_yaml_string(
            yaml_string=rank_and_feedback, expected_keys=expected_keys, cleanup=True
        )

        return pd.DataFrame(
            {
                "Name": rank_and_feedback_dict["name"],
                "Score": rank_and_feedback_dict["score"],
                "Feedback": rank_and_feedback_dict["feedback"],
            }
        )

    def process_submission(
        self,
        video_path: str,
        resume_path: str,
        interview_questions: str,
        job_title: str,
        job_requirements: str,
    ) -> ProcessingResult:
        try:
            # Validate inputs
            error_message = self.validate_inputs(
                video_path, resume_path, interview_questions, job_requirements
            )
            if error_message:
                return ProcessingResult(error_message=error_message)

            # Process inputs
            video_transcript = self.process_video(video_path)
            emotion_analysis = self.analyze_emotions(video_path)
            resume_analysis = self.process_resume(resume_path)

            feedback_list = self.get_feedback(
                interview_questions,
                job_requirements,
                emotion_analysis,
                video_transcript,
                resume_analysis,
            )

            # Update feedback database
            self.candidate_feedback = pd.concat(
                [self.candidate_feedback, feedback_list], ignore_index=True
            )

            # TODO: For testing purposes
            # job_title = "LLM Engineer"
            # interview_questions = INTERVIEW_QUESTION
            # job_requirements = JOB_REQUIREMENTS
            # self.candidate_feedback = pd.DataFrame(
            #     {
            #         "Name": ["Goh Yi Xian"] * 4,
            #         "Score": [50, 50, 50, 50],
            #         "Feedback": [
            #             "The interviewee's technical skills align partially with the job requirements, showcasing proficiency in deep learning frameworks like PyTorch and TensorFlow. However, there is a lack of experience in training and fine-tuning transformer-based models and working with MLOps tools for deployment.",
            #             "The educational background meets the criteria with a Bachelor's degree in Computer Science, but the lack of a Ph.D. and limited industry experience may hinder full alignment with the role.",
            #             "The interview performance indicates a need for improvement in problem-solving skills, confidence, and engagement. The response lacked clarity, relevance, and demonstrated understanding of the key aspects of the job requirements.",
            #             "Overall, while there are some matching skills and experiences, the interviewee falls short in demonstrating a comprehensive fit for the LLM Engineer position. Further development in technical expertise, problem-solving abilities, and communication skills is recommended.",
            #         ],
            #     }
            # )

            write_user_data(
                self.candidate_feedback["Name"].iloc[0],
                self.candidate_feedback["Score"].iloc[0],
                interview_questions,
                job_title,
                job_requirements,
                self.candidate_feedback["Feedback"].tolist(),
            )

            feedback_md = self.format_feedback_to_markdown(self.candidate_feedback)

            return ProcessingResult(
                candidate_name=self.candidate_feedback["Name"].iloc[0],
                candidate_score=self.candidate_feedback["Score"].iloc[0],
                candidate_feedbacks=self.candidate_feedback["Feedback"].tolist(),
                feedback_md=feedback_md,
                interview_question=interview_questions,
                job_requirements=job_requirements,
            )

        except Exception as e:
            self.logger.error(f"Error in process_submission: {str(e)}")
            return ProcessingResult(
                error_message=f"An error occurred during processing: {str(e)}"
            )

    def save_report(
        self,
        candidate_name,
        candidate_score,
        candidate_feedback,
        interview_question,
        job_requirements,
    ) -> Optional[str]:
        try:
            if self.candidate_feedback.empty:
                return None

            doc = Document()
            doc.add_heading(f"Interview Analysis Report - {candidate_name}", 0)
            doc.add_heading("Interview Questions", 1)
            doc.add_paragraph(interview_question)
            doc.add_heading("Job Requirements", 1)
            doc.add_paragraph(job_requirements)
            doc.add_heading("Overall Score", 1)
            paragraph = doc.add_paragraph()
            paragraph.add_run(f"{candidate_score}/100").bold = True
            doc.add_heading("Detailed Feedback", 1)

            for feedback in candidate_feedback:
                doc.add_paragraph(f"• {feedback}")

            doc.save(str(OUTPUT_REPORT_FILE_EMPTY))
            return str(OUTPUT_REPORT_FILE_EMPTY)

        except Exception as e:
            self.logger.error(f"Error saving report: {str(e)}")
            return None

    def create_interface(self) -> gr.Blocks:

        theme = gr.themes.Ocean(
            primary_hue="pink",
            secondary_hue="rose",
            font="Chalkboard",
        )

        with gr.Blocks(title="HR Interview Analysis System", theme=theme) as demo:
            gr.Markdown("# HR Interview Analysis System")

            with gr.Row():
                with gr.Column():
                    video_input = gr.Video(label="Upload Interview Video", format="mp4")
                    resume_input = gr.File(
                        label="Upload Resume (PDF)", file_types=[".pdf"]
                    )

            with gr.Row():
                question_input = gr.Textbox(
                    label="Interview Questions",
                    lines=5,
                    placeholder="Enter the interview questions here...",
                )

            with gr.Row():
                job_title_input = gr.Textbox(
                    label="Job Title",
                    lines=5,
                    placeholder="Enter the job title here...",
                )
                requirements_input = gr.Textbox(
                    label="Job Requirements",
                    lines=5,
                    placeholder="Enter the job requirements here...",
                )

            submit_button = gr.Button("Analyze Interview", variant="primary")

            # Error message display
            error_output = gr.Markdown(visible=False)

            with gr.Tabs():
                with gr.Tab("Analysis Results"):
                    feedback_output_md = gr.Markdown(
                        label="Candidate Assessment",
                        value="No assessment available yet.",
                    )

                    save_button = gr.Button("Generate Report", variant="secondary")
                    report_output = gr.File(label="Download Report")

                with gr.Tab("Candidates List"):
                    candidates_df = gr.Dataframe(
                        headers=[
                            "Name",
                            "Job Title",
                            "Interview Question",
                            "Score",
                            "Feedback",
                        ],
                        datatype=["str", "str", "str", "int", "str"],
                        row_count=(0, "dynamic"),
                        col_count=(5, "fixed"),
                        value=read_all_users(),  # Load initial data
                        interactive=True,
                        wrap=True,
                    )

                    refresh_button = gr.Button("Refresh Candidates List")

                    refresh_button.click(
                        fn=lambda: read_all_users(),  # Reload the candidates data
                        inputs=[],
                        outputs=[candidates_df],
                    )

            candidate_name_state = gr.State()
            candidate_score_state = gr.State()
            candidate_feedbacks_state = gr.State()
            interview_question_state = gr.State()
            job_requirements_state = gr.State()

            # Event handlers
            submit_button.click(
                fn=lambda video, resume, questions, job_title, requirements: (
                    lambda result: (
                        result.candidate_name,
                        result.candidate_score,
                        result.candidate_feedbacks,
                        result.feedback_md,
                        result.interview_question,
                        result.job_requirements,
                        result.error_message,
                    )
                )(
                    self.process_submission(
                        video, resume, questions, job_title, requirements
                    )
                ),
                inputs=[
                    video_input,
                    resume_input,
                    question_input,
                    job_title_input,
                    requirements_input,
                ],
                outputs=[
                    candidate_name_state,
                    candidate_score_state,
                    candidate_feedbacks_state,
                    feedback_output_md,
                    interview_question_state,
                    job_requirements_state,
                    error_output,
                ],
            )

            save_button.click(
                fn=self.save_report,
                inputs=[
                    candidate_name_state,
                    candidate_score_state,
                    candidate_feedbacks_state,
                    interview_question_state,
                    job_requirements_state,
                ],
                outputs=[report_output],
            )

        return demo


def launch_app():
    app = GradioInterface()
    interface = app.create_interface()
    interface.launch(server_name="0.0.0.0", server_port=7860, share=True, debug=True)


if __name__ == "__main__":
    launch_app()
